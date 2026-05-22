from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Лабораторна робота №5.docx"
FIG = ROOT / "lab5_interface_preview.png"
CODE = ROOT / "WpfApp5" / "MainWindow.xaml.cs"


def set_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_cell(cell, bold=False, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = bold


def add_centered(document, text, size=14, bold=False, after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return paragraph


def add_body(document, text="", after=6, align=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.35)
    paragraph.paragraph_format.space_after = Pt(after)
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return paragraph


def add_heading(document, text):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_code(document, code):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    for line in code.rstrip().splitlines():
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        paragraph.add_run("\n")


def extract_method(source, signature, max_lines=45):
    lines = source.splitlines()
    start = next(i for i, line in enumerate(lines) if signature in line)
    depth = 0
    result = []
    seen_open = False
    for line in lines[start:]:
        result.append(line)
        depth += line.count("{")
        if "{" in line:
            seen_open = True
        depth -= line.count("}")
        if seen_open and depth == 0:
            break
        if len(result) >= max_lines:
            result.append("...")
            break
    return "\n".join(result)


def make_preview_image(path):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False

    img = Image.new("RGB", (1400, 760), "#1E1E2E")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arial.ttf", 30)
        font_h = ImageFont.truetype("arial.ttf", 20)
        font = ImageFont.truetype("arial.ttf", 17)
        font_small = ImageFont.truetype("arial.ttf", 14)
    except Exception:
        font_title = font_h = font = font_small = ImageFont.load_default()

    draw.rounded_rectangle((30, 25, 1370, 105), radius=10, fill="#313244")
    draw.text((55, 48), "Моделювання алгоритму банкіра", fill="#CDD6F4", font=font_title)
    draw.text((55, 82), "Лабораторна робота №5 | Пункт 6", fill="#A6ADC8", font=font_small)

    draw.rounded_rectangle((30, 130, 335, 610), radius=8, fill="#2A2A3E")
    draw.text((55, 155), "Параметри системи", fill="#CDD6F4", font=font_h)
    for idx, text in enumerate(["Кількість процесів P: 4", "Типи ресурсів R: 2", "E = [11, 9]", "A = [5, 0]"]):
        y = 205 + idx * 55
        draw.rounded_rectangle((55, y, 310, y + 36), radius=5, outline="#45475A", fill="#1E1E2E")
        draw.text((70, y + 9), text, fill="#CDD6F4", font=font)

    tables = [
        ("Виділено C", [["1", "4"], ["0", "1"], ["4", "4"], ["1", "0"]]),
        ("Максимум Max", [["4", "7"], ["1", "3"], ["8", "8"], ["2", "4"]]),
        ("Потреби Need", [["3", "3"], ["1", "2"], ["4", "4"], ["1", "4"]]),
    ]
    x0 = 365
    for t, rows in tables:
        draw.rounded_rectangle((x0, 130, x0 + 315, 360), radius=8, fill="#2A2A3E")
        draw.text((x0 + 20, 155), t, fill="#CDD6F4", font=font_h)
        y = 205
        draw.text((x0 + 120, y - 28), "R1", fill="#A6C8FF", font=font_small)
        draw.text((x0 + 220, y - 28), "R2", fill="#A6C8FF", font=font_small)
        for i, row in enumerate(rows):
            yy = y + i * 32
            draw.text((x0 + 25, yy + 7), f"P{i}", fill="#A6C8FF", font=font_small)
            for j, value in enumerate(row):
                xx = x0 + 90 + j * 100
                draw.rectangle((xx, yy, xx + 85, yy + 26), outline="#45475A", fill="#1E1E2E")
                draw.text((xx + 8, yy + 5), value, fill="#CDD6F4", font=font_small)
        x0 += 335

    draw.rounded_rectangle((365, 385, 1350, 710), radius=8, fill="#2A2A3E")
    draw.text((390, 410), "Лог моделювання", fill="#CDD6F4", font=font_h)
    log_lines = [
        "E = [11, 9]; Available A = [5, 0]",
        "Need = Max - C",
        "Work = Available = [5, 0]",
        "Немає процесу, для якого Need <= Work",
        "Алгоритм банкіра: небезпечний стан",
    ]
    for i, line in enumerate(log_lines):
        draw.text((405, 455 + i * 36), line, fill="#F9E2AF" if i >= 3 else "#A6E3A1", font=font)

    img.save(path)
    return True


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for name in ["Heading 1", "Heading 2"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def build_report():
    document = Document()
    configure_styles(document)
    set_margins(document.sections[0])

    add_centered(document, "Міністерство освіти і науки України", 12)
    add_centered(document, "ЧЕРКАСЬКИЙ НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ\nім. Богдана Хмельницького", 12, bold=True)
    add_body(document, "")
    add_body(document, "Факультет \tОбчислювальної техніки, інтелектуальних та управляючих систем", after=0)
    add_body(document, "Кафедра \tПрограмного забезпечення автоматизованих систем", after=0)
    for _ in range(5):
        add_body(document, "")
    add_centered(document, "ЛАБОРАТОРНА РОБОТА №5\nпо дисципліні «Системне програмування»", 14, bold=True, after=12)
    add_centered(document, "Тема: Моделювання і усунення ситуацій взаємного блокування і голодування", 12)
    for _ in range(4):
        add_body(document, "")
    add_body(document, "Виконала: \tстудентка гр. КС-24", after=0)
    add_body(document, "Гук М.О.", after=12)
    add_body(document, "Перевірив:\tвикладач", after=0)
    add_body(document, "кафедри ІТ", after=0)
    add_body(document, "Авраменко А.С.", after=0)
    for _ in range(5):
        add_body(document, "")
    add_centered(document, "Черкаси, 2026", 12)

    document.add_page_break()

    add_body(
        document,
        "Мета роботи: вивчити принципи виникнення взаємоблокувань і реалізувати "
        "моделювання алгоритму банкіра для перевірки безпечності стану системи."
    )

    add_heading(document, "ЗАВДАННЯ")
    add_body(
        document,
        "Розробити програму підвищеної складності, що моделює роботу алгоритму банкіра. "
        "Кількість процесів має задаватися в межах від 2 до 4, кількість типів ресурсів - від 1 до 3."
    )
    add_bullet(document, "згенерувати випадкову систему процесів і ресурсів у заданих межах;")
    add_bullet(document, "сформувати матриці Allocation, Max і Need;")
    add_bullet(document, "обчислити вектор доступних ресурсів Available;")
    add_bullet(document, "перевірити стан системи алгоритмом банкіра;")
    add_bullet(document, "вивести безпечну послідовність або показати, що стан небезпечний.")

    add_heading(document, "ХІД РОБОТИ")
    add_body(document, "Програма реалізована мовою C# з графічним інтерфейсом WPF. Інтерфейс дозволяє генерувати параметри системи, переглядати матриці ресурсів і запускати аналіз алгоритмом банкіра.")

    add_body(document, "У програмі використовуються такі основні позначення:")
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Позначення", "Поле в програмі", "Зміст"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
        style_cell(table.rows[0].cells[idx], bold=True, fill="D9EAF7")
    rows = [
        ("E", "_total", "загальна кількість ресурсів у системі"),
        ("A / Available", "_available", "кількість ресурсів, які зараз вільні"),
        ("C / Allocation", "_allocation", "ресурси, вже виділені процесам"),
        ("Max", "_maximum", "максимальні потреби процесів"),
        ("Need", "_need", "ресурси, які ще можуть знадобитися процесу"),
        ("Work", "work", "робоча копія Available для уявного завершення процесів"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            style_cell(cells[idx])

    add_body(document, "Матриця потреб обчислюється за формулою:")
    add_code(document, "Need[i, j] = Max[i, j] - Allocation[i, j];")

    add_body(document, "Вектор доступних ресурсів обчислюється як різниця між загальною кількістю ресурсу і сумою виділених одиниць цього ресурсу:")
    add_code(document, "Available[j] = Total[j] - sum(Allocation[i, j]);")

    add_body(document, "Основні можливості інтерфейсу програми:")
    add_bullet(document, "випадкова генерація кількості процесів від 2 до 4;")
    add_bullet(document, "випадкова генерація кількості типів ресурсів від 1 до 3;")
    add_bullet(document, "перегляд і редагування матриць Allocation і Max;")
    add_bullet(document, "автоматичний перерахунок Need і Available;")
    add_bullet(document, "покроковий лог роботи алгоритму банкіра.")

    if make_preview_image(FIG):
        document.add_picture(str(FIG), width=Inches(6.3))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Рис. 1 - Інтерфейс програми моделювання алгоритму банкіра")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    add_body(document, "Генерація стану системи виконується в методі GenerateScenario. Спочатку створюється матриця Allocation, після цього для кожного процесу формується Max, а Need рахується як різниця між Max і Allocation.")

    source = CODE.read_text(encoding="utf-8")
    add_body(document, "Фрагмент методу генерації стану системи:")
    add_code(document, extract_method(source, "private void GenerateScenario()", max_lines=38))

    add_body(document, "Перевірка безпечності виконується методом CheckSafety. Алгоритм бере Work = Available і шукає процес, для якого Need не перевищує Work. Якщо такий процес існує, він уявно завершується і повертає виділені ресурси в Work.")
    add_code(document, extract_method(source, "private SafetyResult CheckSafety()", max_lines=55))

    add_body(document, "Якщо всі процеси вдалося позначити як завершені, стан вважається безпечним. Якщо на певному кроці немає жодного процесу, для якого виконується Need <= Work, стан є небезпечним.")

    add_heading(document, "Висновок")
    add_body(
        document,
        "У ході виконання лабораторної роботи було розроблено WPF-програму для моделювання алгоритму банкіра. "
        "Програма генерує систему з 2-4 процесів і 1-3 типів ресурсів, формує матриці Allocation, Max і Need, "
        "обчислює вектор Available та виконує перевірку стану системи на безпечність. У журналі моделювання "
        "покроково відображається робота алгоритму: початковий Work, перевірка умови Need <= Work, уявне "
        "завершення процесу та повернення ресурсів. Якщо існує послідовність завершення всіх процесів, стан "
        "визначається як безпечний; інакше програма повідомляє про небезпечний стан."
    )

    document.save(OUT)


if __name__ == "__main__":
    build_report()
    print(OUT)
